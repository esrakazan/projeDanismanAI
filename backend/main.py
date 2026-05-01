from fastapi import FastAPI, Depends, HTTPException, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel, EmailStr
from typing import Optional, List
import io
 
from database import get_db, User, Conversation, Message, create_tables
from auth import hash_password, verify_password, create_access_token, get_current_user
from skill_router import detect_skill
from hf_client import call_hf_endpoint, build_prompt
 
app = FastAPI(title="ProjeDanışmanAI API")
 
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
 
 
@app.on_event("startup")
def startup():
    create_tables()
 
 
# ---------- SCHEMAS ----------
 
class RegisterRequest(BaseModel):
    email: EmailStr
    password: str
 
class LoginRequest(BaseModel):
    email: EmailStr
    password: str
 
class ChatRequest(BaseModel):
    conversation_id: int
    message: str
    skill: Optional[str] = "otomatik"
 
class ConversationCreate(BaseModel):
    title: Optional[str] = "Yeni Sohbet"
 
 
# ---------- AUTH ----------
 
@app.post("/auth/register")
def register(req: RegisterRequest, db: Session = Depends(get_db)):
    if db.query(User).filter(User.email == req.email).first():
        raise HTTPException(status_code=400, detail="Bu email zaten kayıtlı")
    user = User(email=req.email, hashed_password=hash_password(req.password))
    db.add(user)
    db.commit()
    db.refresh(user)
    return {"message": "Kayıt başarılı", "email": user.email}
 
 
@app.post("/auth/login")
def login(req: LoginRequest, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == req.email).first()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(status_code=401, detail="Email veya şifre hatalı")
    token = create_access_token({"sub": user.email})
    return {"access_token": token, "token_type": "bearer"}
 
 
@app.get("/auth/me")
def me(current_user: User = Depends(get_current_user)):
    return {"id": current_user.id, "email": current_user.email}
 
 
# ---------- CONVERSATIONS ----------
 
@app.post("/conversations")
def create_conversation(req: ConversationCreate, db: Session = Depends(get_db),
                        current_user: User = Depends(get_current_user)):
    conv = Conversation(user_id=current_user.id, title=req.title)
    db.add(conv)
    db.commit()
    db.refresh(conv)
    return {"id": conv.id, "title": conv.title, "created_at": conv.created_at}
 
 
@app.get("/conversations")
def list_conversations(db: Session = Depends(get_db),
                       current_user: User = Depends(get_current_user)):
    convs = db.query(Conversation).filter(
        Conversation.user_id == current_user.id
    ).order_by(Conversation.created_at.desc()).all()
    return [{"id": c.id, "title": c.title, "created_at": c.created_at} for c in convs]
 
 
@app.get("/conversations/{conv_id}/messages")
def get_messages(conv_id: int, db: Session = Depends(get_db),
                 current_user: User = Depends(get_current_user)):
    conv = db.query(Conversation).filter(
        Conversation.id == conv_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Sohbet bulunamadı")
    msgs = db.query(Message).filter(
        Message.conversation_id == conv_id
    ).order_by(Message.created_at).all()
    return [{"role": m.role, "content": m.content, "skill_used": m.skill_used,
             "created_at": m.created_at} for m in msgs]
 
 
@app.delete("/conversations/{conv_id}")
def delete_conversation(conv_id: int, db: Session = Depends(get_db),
                        current_user: User = Depends(get_current_user)):
    conv = db.query(Conversation).filter(
        Conversation.id == conv_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Sohbet bulunamadı")
    db.query(Message).filter(Message.conversation_id == conv_id).delete()
    db.delete(conv)
    db.commit()
    return {"message": "Sohbet silindi"}
 
 
# ---------- CHAT ----------
 
@app.post("/chat")
def chat(req: ChatRequest, db: Session = Depends(get_db),
         current_user: User = Depends(get_current_user)):
    conv = db.query(Conversation).filter(
        Conversation.id == req.conversation_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Sohbet bulunamadı")
 
    skill_name = req.skill if req.skill != "otomatik" else detect_skill(req.message)
 
    recent = db.query(Message).filter(
        Message.conversation_id == req.conversation_id
    ).order_by(Message.created_at.desc()).limit(8).all()
    history = [{"role": m.role, "content": m.content} for m in reversed(recent)]
 
    prompt = build_prompt(skill_name, history, req.message)
    response_text = call_hf_endpoint(prompt)
 
    db.add(Message(conversation_id=req.conversation_id, role="user",
                   content=req.message, skill_used=skill_name))
    db.add(Message(conversation_id=req.conversation_id, role="assistant",
                   content=response_text, skill_used=skill_name))
 
    if conv.title == "Yeni Sohbet":
        conv.title = req.message[:40] + ("..." if len(req.message) > 40 else "")
 
    db.commit()
 
    return {"response": response_text, "skill_used": skill_name}
 
 
@app.post("/chat/word")
def chat_word(req: ChatRequest, db: Session = Depends(get_db),
              current_user: User = Depends(get_current_user)):
    # Sohbetin bu kullanıcıya ait olduğunu doğrula
    conv = db.query(Conversation).filter(
        Conversation.id == req.conversation_id,
        Conversation.user_id == current_user.id
    ).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Sohbet bulunamadı")
 
    # En son assistant mesajını al, tekrar model çağırma
    last_msg = db.query(Message).filter(
        Message.conversation_id == req.conversation_id,
        Message.role == "assistant"
    ).order_by(Message.created_at.desc()).first()
 
    if not last_msg:
        raise HTTPException(status_code=404, detail="İndirilecek mesaj bulunamadı")
 
    response_text = last_msg.content
 
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("ProjeDanışmanAI Raporu", 0)
        for line in response_text.split("\n"):
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.strip():
                doc.add_paragraph(line)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=rapor.docx"}
        )
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx kurulu değil")
 
 
@app.get("/skills")
def list_skills():
    return [
        {"value": "otomatik", "label": "Otomatik Tespit"},
        {"value": "project_idea_refinement", "label": "Proje Fikri Netleştirme"},
        {"value": "report_section_writer", "label": "Rapor Bölümü Yazma"},
        {"value": "tubitak_application_guidance", "label": "TÜBİTAK Başvuru Rehberliği"},
        {"value": "teknofest_ktr_ptr_guidance", "label": "TEKNOFEST KTR/PTR"},
        {"value": "feasibility_and_risk_check", "label": "Uygulanabilirlik ve Risk"},
        {"value": "title_abstract_generator", "label": "Başlık ve Özet"},
        {"value": "presentation_and_jury_preparation", "label": "Sunum ve Jüri Hazırlığı"},
    ]
 
 
@app.get("/health")
def health():
    return {"status": "ok"}